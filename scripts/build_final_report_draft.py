"""Build an editable course-report draft from the supplied 2026 template."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"D:\AI+ Project Report - Template 2026.docx")
OUTPUT = ROOT / "deliverables" / "ProxyGap_Final_Report_Draft_20260820.docx"
FIGURES = ROOT / "docs" / "figures" / "v4_multiobjective_final"


OVERVIEW = [
    "Reward maximisation alone does not guarantee that a quadruped follows a requested direction, remains safe, or completes a mission. This project therefore treated reward as a training signal rather than as the definition of success. It used a locally reproduced Gymnasium Ant-v5 and proximal policy optimisation (PPO) configuration as the matched numerical baseline, then developed two linked stages. Stage 1 (Project V2) investigated directed flat-ground locomotion through target-speed, body-orientation, action-smoothness and contact diagnostics. Stage 2 (Project V3) moved the task to an 80 m × 80 m continuous heightfield and separated contact support, slope capability, turning, global route selection and human time–energy preferences.",
    "The final system was hierarchical. A frozen known-map planner supplied local waypoints to an archived bidirectional low-level expert, while a four-pair PAIR0 contact contract reduced a heightfield contact-margin artefact. Fifteen route and speed candidates first had to satisfy arrival and safety gates. Three preference profiles then ranked only valid candidates by normalised completion time and positive mechanical-work proxy. Time-priority and balanced preferences selected one route contract; energy-priority selected a second. Across three new reset seeds per contract, all six formal episodes reached the goal and completed the required two-second dwell, with no falls, torso-ground contacts, sustained non-foot contacts or duration-corrected sustained-slip events.",
    "These results are deliberately bounded. They concern one previously inspected frozen map, two unique route contracts and three reset seeds, rather than unseen-map generalisation or independent training-seed robustness. Positive mechanical work is a simulation proxy, not battery energy. Representative successful episodes still contained 9.25–10.02% complete control intervals with no foot contact, so neither a biologically natural gait nor continuous ground support is claimed. The principal research contribution is consequently an auditable decomposition: support, turning and planning were different failure layers, and reliable completion required a system architecture rather than one further reward term.",
]

CONTEXT = [
    "The defensible baseline combines an environment specification, an algorithmic source and a matched local implementation. Farama Foundation's Ant page defines the current Gymnasium Ant-v5 environment and its eight torque actions, but it is documentation rather than a research paper (Farama Foundation, n.d.). PPO supplies the algorithmic provenance: its clipped surrogate objective supports stable minibatch policy updates in continuous-control tasks (Schulman et al., 2017). Earlier generalised advantage estimation work includes high-dimensional locomotion and provides historical context for Ant-style control (Schulman et al., 2016). Numerical claims in this report use local matched conditions because published studies differ in robot morphology, observations, rewards, simulator versions and training budgets.",
    "Quadruped literature nevertheless informs the design. Lee et al. (2020) demonstrate robust learned locomotion over challenging terrain, while Miki et al. (2022) show how local exteroceptive information can support anticipatory control in the wild. Aractingi et al. (2023) combine command tracking, body orientation, smoothness and effort considerations for a physical quadruped. Fu et al. (2022) show that energy objectives can produce gait transitions, providing a methodological anchor for measuring mechanical work and contact timing. Their scores are not directly comparable with this Ant system, so they are used to justify metrics rather than as imported numerical baselines.",
    "Reward design also creates an evaluation risk. Potential-based shaping can preserve an optimal policy under specific assumptions (Ng, Harada and Russell, 1999), but most practical locomotion shaping terms do not meet that guarantee. Misspecified rewards can therefore favour behaviours that score well while violating the intended task (Pan, Bhatia and Steinhardt, 2022; Skalse et al., 2022). Stepwise action exploration may additionally create jerky robot motion (Raffin, Kober and Stulp, 2022). These concerns motivated separate task-validity, safety, gait, navigation and work-proxy measurements. Stage 1 compares interventions with the matched Ant-v5/PPO baseline; Stage 2 uses sequential ablations and the direct-goal system as fairer comparators than asking an unmodified Ant controller to solve a task for which it receives neither map nor waypoint information.",
]

ROLES = [
    "[Replace before submission.] Member A (Name in Pinyin) coordinated experiment protocols and reproducibility; Member B implemented and tested locomotion and contact interventions; Member C analysed results, figures and literature; Member D integrated the planner, videos, report and presentation. The final version must replace these placeholders with the actual contribution record, identify shared work honestly and ensure that every member can explain the methods and evidence used in the presentation.",
]

STAGE1 = [
    "The legacy V1 experiments showed that a generic forward-return objective did not define the team's intended behaviour. Stage 1 therefore introduced bounded target-speed and direction tracking, torso vertical and angular terms, a signed-pitch diagnostic, action-rate and saturation measurements, and foot-landing/contact terms. The Ant body and eight-dimensional torque action space were retained. An external slew limiter was recorded as a controller intervention rather than described as learned gait intelligence. Independent measures included net displacement, heading error, path efficiency, torso attitude, action roughness, contact order and no-floor exposure.",
    "Two development comparisons illustrate partial improvement. In one paired setting, normalised action roughness decreased from 0.0139 to 0.00985, mean speed increased from 0.844 to 0.918 m s−1 and path efficiency increased from 0.809 to 0.857. In a second, mean take-offs decreased from 21.1 to 3.77 and no-floor fraction from 0.526 to 0.465, while mean speed decreased from 0.961 to 0.931 m s−1. These are development diagnostics, not independent training replications. They support the cautious conclusion that motion looked more coordinated and selected proxies improved. Duty factor, phase and contact ordering were not frozen across speeds as biological validation metrics; consequently, “natural gait” remains a visual description rather than a scientific finding.",
]

CONTACT = [
    "Stage 2 first added a 13-dimensional local terrain preview to the 122-dimensional locomotion observation. Nine relative-height samples, a target-frame terrain normal and signed forward slope produced a 135-dimensional policy input without giving the low-level policy the complete map. A plane-versus-heightfield diagnosis then isolated a major contact discrepancy. With default positive geom margins, heightfield–capsule collision semantics created many additional contacts and high zero-foot exposure. The mechanism was treated as a frozen simulator/contact interaction, not labelled a MuJoCo bug.",
    "The PAIR0 contract retained all geom margins and introduced four explicit floor-to-distal-foot contact pairs with zero pair margin and gap. In the held-out diagnostic, full-control zero-foot exposure decreased from approximately 24.04% to 3.21%, mean supporting feet increased from 0.353 to 1.344 and the best-progress ratio was 1.051. Safety remained strict: finite state, no fall, no torso-ground, no sustained non-foot contact, non-zero force-qualified support denominator, and no duration-corrected sustained-slip event. PAIR0 improved support under the frozen model, but support alone did not provide route choice or reliable steering.",
]

DIAGNOSIS = [
    "Several negative results were retained because they changed the system design. Reducing commanded speed to 0.20–0.40 m s−1 did not satisfy the predeclared full-interval support gate. A grouped terrain-frame reward switch reduced progress and introduced an additional fall, so it was rejected rather than tuned retrospectively. A phase-aware crawl smoke test improved some observations but did not reach the formal promotion threshold. The slope grid established only a tested capability boundary: uphill episodes passed contiguously through 12°, with 16° the first tested failure; downhill results were non-monotonic. Across 55 slope episodes there were no falls, torso-ground, sustained non-foot or corrected sustained-slip events, but 12° is not a physical maximum.",
    "Turning remained a separate limitation. Balanced left/right command exposure preserved slope safety but failed the signed turn-tracking gate. The formal direct-goal PAIR0 policy then demonstrated why local support was insufficient: in 600 s it never entered the goal region, achieving only 14.51 m best progress and 12.53 m net progress along a 92.41 m executed path. It received local terrain and a target direction, but no full-map route, and its bidirectional response was unreliable. This negative result justified a high-level planner; it was not hidden or replaced by a favourable seed.",
]

HIERARCHY = [
    "A read-only checkpoint screen identified an archived canonical-frame V4 expert as the only candidate with both left- and right-turn response. When combined with PAIR0, it completed flat, ±8° and +12° safety screens without falls or sustained corrected-slip events. A naïve action blend between V4 and the final PAIR0 policy fell after 15.3 s and was rejected. The retained architecture instead kept responsibilities explicit: a known-map planner searched the frozen 1025 × 1025 heightfield, a waypoint follower supplied a local target with a 3 m lookahead, and the V4 policy produced the eight joint torques. The route screen used a 16° discrete corridor-slope proxy. It is therefore a known-map system integration, not evidence that the low-level policy learned global planning.",
    "Fifteen feasible route and speed candidates were evaluated. Arrival required entry within 1.5 m followed by forty consecutive 0.05 s control steps inside a 2 m circle. A candidate also required finite states, no fall, no torso-ground, no sustained non-foot contact and no duration-corrected sustained-slip event. Only candidates passing these lexicographic gates entered preference ranking. The objective J = wT(T/Tmin) + wE(W+/W+min) used declared weights: 0.8/0.2 for time priority, 0.5/0.5 for balanced and 0.2/0.8 for energy priority. Time and balanced preferences selected the same route; energy priority selected a second. This establishes near-optimality within the evaluated bank, not a continuous global optimum.",
    "Formal evaluation used three new hash-derived reset seeds per route contract. The time/balanced contract completed 3/3 episodes with mean time 264.55 s, mean positive-work proxy 55.65 kJ and mean path length 153.23 m. The energy contract also completed 3/3, with corresponding means of 259.37 s, 55.13 kJ and 152.14 m. Reset variability reversed the development-time ordering—the energy-selected route was slightly faster in the formal mean—so all seed points are reported rather than claiming a deterministic trade-off. Each representative video was reproduced from its formal control trace with zero state and five-substep contact mismatches, then decoded frame by frame.",
    "Reproducibility was handled as part of the experimental method rather than as post-processing. Formal attempts used fail-if-nonempty output roots, frozen configuration files, hash-derived seeds and SHA-256 inventories for the checkpoint, map assets, routes, traces and videos. The final runner verified that the checkpoint timestep and hashes were unchanged and wrote no replacement model. Tests covered route scoring, preference selection, manifest closure and exact replay; the repository-level run reported 408 passes and three skips. These controls do not remove scientific uncertainty, but they prevent a favourable video from becoming detached from the evaluated trajectory. They also make negative results interpretable: a failed direct-goal episode, rejected action blend or failed turning gate remains associated with its actual code, seed and contact contract instead of being silently overwritten.",
]

CONTRIBUTIONS = [
    "The project contributes an auditable two-stage evaluation framework that distinguishes optimised reward from independent human-intent and safety measures. It provides a contact-focused diagnosis of heightfield–capsule margin behaviour, a predeclared PAIR0 contract, explicit slope and turn boundaries, and a lexicographic mission evaluation in which validity and safety precede preference. The final hierarchical known-map system achieved six completions from six formal episodes, while preserving rejected interventions and exact replay evidence. This is more informative than presenting only the highest return: each retained failure identified whether the missing capability concerned contact, steering or planning.",
    "The evidence has important limitations. All final runs use one known, previously inspected map and only three reset seeds; independent training-seed replication and unseen start–goal/map splits remain future work. The archived V4 expert was selected during development. Fifteen candidates do not cover the continuous route and controller space. Positive and absolute mechanical work omit motor efficiency, electronics and battery dynamics, so electrical energy and mission energy are not established. No real robot, sensor noise or actuator uncertainty was tested. Successful representative episodes still contained 9.25–10.02% full-control zero-foot intervals; zero sustained-slip events does not mean zero transient slip candidates or uninterrupted support. The method therefore establishes reproducible known-map completion under declared constraints, not universal terrain locomotion, biological natural gait or global optimality.",
    "The next defensible study would freeze this architecture and evaluate independent training seeds, unseen maps and unseen start–goal pairs. A calibrated actuator or battery model should replace the work proxy if energy is to become a primary claim. Gait evaluation should predeclare per-foot duty factors, phase consistency, contact order and airborne limits across commanded speeds. These extensions are preferable to further unconstrained reward scanning because they directly test the remaining evidence gaps.",
]

REFERENCES = [
    "Aractingi, M., Despré, R., Righetti, L. and Wolf, P. (2023) ‘Controlling the Solo12 quadruped robot with deep reinforcement learning’, Scientific Reports, 13, 11945. doi: 10.1038/s41598-023-38259-7.",
    "Farama Foundation (n.d.) Ant. Gymnasium Documentation. Available at: https://gymnasium.farama.org/environments/mujoco/ant/ (Accessed: 20 August 2026).",
    "Fu, Z., Kumar, A., Malik, J. and Pathak, D. (2022) ‘Minimizing energy consumption leads to the emergence of gaits in legged robots’, Proceedings of Machine Learning Research, 164, pp. 928–937.",
    "Lee, J. et al. (2020) ‘Learning quadrupedal locomotion over challenging terrain’, Science Robotics, 5(47), eabc5986.",
    "Miki, T. et al. (2022) ‘Learning robust perceptive locomotion for quadrupedal robots in the wild’, Science Robotics, 7(62), eabk2822.",
    "Ng, A.Y., Harada, D. and Russell, S.J. (1999) ‘Policy invariance under reward transformations: Theory and application to reward shaping’, Proceedings of ICML, pp. 278–287.",
    "Pan, A., Bhatia, K. and Steinhardt, J. (2022) ‘The effects of reward misspecification: Mapping and mitigating misaligned models’, International Conference on Learning Representations.",
    "Raffin, A., Kober, J. and Stulp, F. (2022) ‘Smooth exploration for robotic reinforcement learning’, Proceedings of Machine Learning Research, 164, pp. 1634–1644.",
    "Schulman, J. et al. (2016) ‘High-dimensional continuous control using generalized advantage estimation’, International Conference on Learning Representations.",
    "Schulman, J. et al. (2017) ‘Proximal policy optimization algorithms’, arXiv:1707.06347.",
    "Skalse, J. et al. (2022) ‘Defining and characterizing reward hacking’, Advances in Neural Information Processing Systems, 35.",
]


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "A8B0B5")
        borders.append(element)


def add_heading(document: Document, text: str, level: int = 1):
    p = document.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_paragraphs(document: Document, texts: list[str]) -> None:
    for text in texts:
        p = document.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.24)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15


def add_figure(document: Document, path: Path, caption: str, width: float = 6.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = document.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = "Caption"


def count_body_words() -> int:
    parts = OVERVIEW + CONTEXT + ROLES + STAGE1 + CONTACT + DIAGNOSIS + HIERARCHY + CONTRIBUTIONS
    return len(re.findall(r"\b[\w–+−/.]+\b", " ".join(parts)))


def main() -> None:
    document = Document(TEMPLATE)
    if "Caption" not in [style.name for style in document.styles]:
        caption_style = document.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.name = "Times New Roman"
        caption_style.font.size = Pt(9)
        caption_style.font.italic = True
    body = document.element.body
    first_table = document.tables[0]._tbl
    section_properties = body.sectPr
    for child in list(body):
        if child is not first_table and child is not section_properties:
            body.remove(child)

    title = document.add_heading("Project Report", level=0)
    body.remove(title._p)
    body.insert(0, title._p)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = document.tables[0]
    set_cell_text(info.cell(0, 0), "Programme Cohort", bold=True)
    set_cell_text(info.cell(0, 1), "Reinforcement Learning & Robotic Automation")
    set_cell_text(info.cell(1, 0), "Course Group", bold=True)
    set_cell_text(info.cell(1, 1), "[Insert course group]")
    set_cell_text(info.cell(1, 2), "Group Name", bold=True)
    set_cell_text(info.cell(1, 3), "[Insert group name]")
    set_cell_text(info.cell(2, 0), "Group Members", bold=True)
    set_cell_text(info.cell(2, 1), "[Insert names and student IDs]")

    add_heading(document, "Project Title")
    p = document.add_paragraph("From Reward Shaping to Hierarchical Terrain Navigation: An Auditable PPO–Ant Study")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)

    add_heading(document, "Project Overview")
    add_paragraphs(document, OVERVIEW)

    add_heading(document, "Research Context")
    add_paragraphs(document, CONTEXT)

    add_heading(document, "Division of Roles and Responsibilities")
    add_paragraphs(document, ROLES)

    add_heading(document, "Challenges and Solutions")
    add_heading(document, "Stage 1: From default reward to intent diagnostics", level=2)
    add_paragraphs(document, STAGE1)
    add_heading(document, "Stage 2A: Terrain observation and contact support", level=2)
    add_paragraphs(document, CONTACT)
    add_heading(document, "Stage 2B: Slope, turning and direct-goal failures", level=2)
    add_paragraphs(document, DIAGNOSIS)
    add_heading(document, "Stage 2C: Hierarchical completion and preference selection", level=2)
    add_paragraphs(document, HIERARCHY)

    add_figure(document, FIGURES / "candidate_time_work_scatter.png", "Figure 1. Fifteen feasible candidates and the two selected route contracts. Mechanical work is a simulation proxy, not battery energy.")
    add_figure(document, FIGURES / "representative_planned_vs_actual_routes.png", "Figure 2. Planned waypoints and executed trajectories for the three representative formal episodes.")

    add_heading(document, "Final Results")
    table = document.add_table(rows=2, cols=7)
    add_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Contract", "Success", "Mean time (s)", "Mean W+ proxy (kJ)", "Mean path (m)", "Sustained slips", "Falls"]
    values = [
        ["time/balanced", "3/3", "264.55", "55.65", "153.23", "0", "0"],
        ["energy", "3/3", "259.37", "55.13", "152.14", "0", "0"],
    ]
    for j, text in enumerate(headers):
        set_cell_text(table.cell(0, j), text, bold=True)
        shade_cell(table.cell(0, j), "DDECEC")
    for row in values:
        cells = table.add_row().cells
        for j, text in enumerate(row):
            set_cell_text(cells[j], text)
    table._tbl.remove(table.rows[1]._tr)
    add_figure(document, FIGURES / "formal_per_seed_outcomes.png", "Figure 3. Completion time, positive-work proxy and path length for all formal reset seeds.")

    add_heading(document, "Contributions and Limitations")
    add_paragraphs(document, CONTRIBUTIONS)

    add_heading(document, "Bibliography")
    for ref in REFERENCES:
        p = document.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)

    add_heading(document, "Evidence Appendix")
    evidence = [
        "Candidate selection manifest: 212b78865714332cf28c9c1dc5dc1b8f7fb74883f03ef1ebee3789161c31f20b.",
        "Formal evaluation manifest: 0bf2817cbdaadc02929da91bae7acb04371ff9cf1ec43e1ab4efa3a8a4a08d83.",
        "Video archive manifest: 60498def5e209959e0ea9fd09629b71bc2e6df292dfb00d2140cf838e1ddc024.",
        "Archived V4 checkpoint: 6a0f6081e6aff4c85201242e53c44b0d057e96167336002da3a6e862fe134b6a.",
        "Verification: 408 tests passed and 3 skipped; representative video state/contact replay mismatch = 0; all video frames decoded.",
        "AI assistance disclosure: Codex supported code execution, evidence auditing, drafting and layout. Group members must verify, edit and understand the final submission and follow the course policy.",
    ]
    for item in evidence:
        p = document.add_paragraph(f"•  {item}")
        p.paragraph_format.left_indent = Inches(0.20)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(3)

    document.add_page_break()
    add_heading(document, "Evaluation Sheet")
    p = document.add_paragraph("This section is to be completed by the Instructor(s). Please do not delete.")
    p.runs[0].italic = True
    mark = document.add_table(rows=1, cols=2)
    add_table_borders(mark)
    set_cell_text(mark.cell(0, 0), "Final Mark", bold=True)
    set_cell_text(mark.cell(0, 1), "")
    comments = document.add_table(rows=2, cols=1)
    add_table_borders(comments)
    set_cell_text(comments.cell(0, 0), "Further Comments", bold=True)
    set_cell_text(comments.cell(1, 0), "\n\n\n\n\n")

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    for style_name in ["Title", "Heading 1", "Heading 2", "Caption"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
    styles["Heading 1"].font.color.rgb = RGBColor(22, 140, 140)
    styles["Heading 2"].font.color.rgb = RGBColor(46, 110, 158)

    for section in document.sections:
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.78)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "ProxyGap final report draft · 20 August 2026"
        footer.runs[0].font.name = "Times New Roman"
        footer.runs[0].font.size = Pt(8)

    core = document.core_properties
    core.title = "From Reward Shaping to Hierarchical Terrain Navigation"
    core.subject = "AI+ Project Report 2026"
    core.author = "ProxyGap project team"
    core.keywords = "PPO, Ant-v5, quadruped, reward shaping, terrain navigation"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"{OUTPUT}\nbody_word_count={count_body_words()}")


if __name__ == "__main__":
    main()
